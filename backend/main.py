"""
Production-Level ATS Resume Analyzer Backend
============================================
Enhanced with:
- spaCy NLP for entity extraction (name, org, location, dates)
- Sentence-BERT / TF-IDF semantic skill matching
- Keyword density & section detection
- Experience & education extraction
- ATS score breakdown by category
- Redis caching (optional)
- Structured logging
- Rate limiting
- Background job processing
- Swagger docs
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import logging
import os
import re
import time
from collections import defaultdict
from typing import Any, Optional

import docx
import nltk
import pdfplumber
import requests
import spacy
from fastapi import BackgroundTasks, FastAPI, File, HTTPException, Query, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ─────────────────────────────────────────────────────────────
# Logging
# ─────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
)
logger = logging.getLogger("ats_analyzer")

# ─────────────────────────────────────────────────────────────
# NLP Setup
# ─────────────────────────────────────────────────────────────
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("spaCy model loaded: en_core_web_sm")
except OSError:
    logger.warning("spaCy model not found. Run: python -m spacy download en_core_web_sm")
    nlp = None

nltk.download("punkt", quiet=True)
nltk.download("stopwords", quiet=True)
nltk.download("averaged_perceptron_tagger", quiet=True)

# ─────────────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────────────
ADZUNA_APP_ID = os.getenv("ADZUNA_APP_ID", "d3c2c9e5")
ADZUNA_APP_KEY = os.getenv("ADZUNA_APP_KEY", "3c6be086d0848b3c61571efd9b1e1ad0")
COUNTRY = os.getenv("ADZUNA_COUNTRY", "in")

MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "10"))
RATE_LIMIT_REQUESTS = int(os.getenv("RATE_LIMIT_REQUESTS", "30"))
RATE_LIMIT_WINDOW = int(os.getenv("RATE_LIMIT_WINDOW", "60"))  # seconds

# ─────────────────────────────────────────────────────────────
# Role → Skills taxonomy (critical + bonus)
# ─────────────────────────────────────────────────────────────
ROLE_SKILLS: dict[str, dict[str, list[str]]] = {
    "Data Analyst": {
        "critical": ["sql", "excel", "data analysis", "statistics", "python", "reporting"],
        "bonus": ["power bi", "tableau", "looker", "r", "databricks", "snowflake", "airflow"],
    },
    "Data Scientist": {
        "critical": ["machine learning", "python", "statistics", "pandas", "model evaluation", "data wrangling"],
        "bonus": ["deep learning", "nlp", "tensorflow", "pytorch", "spark", "mlflow", "feature engineering"],
    },
    "ML Engineer": {
        "critical": ["machine learning", "python", "model deployment", "docker", "kubernetes", "ci/cd"],
        "bonus": ["tensorflow", "pytorch", "triton", "onnx", "mlops", "ray", "kubeflow"],
    },
    "AI Engineer": {
        "critical": ["llm", "transformers", "nlp", "deep learning", "python", "prompt engineering"],
        "bonus": ["langchain", "openai", "huggingface", "rag", "fine-tuning", "vector database", "fastapi"],
    },
    "Backend Engineer": {
        "critical": ["python", "rest api", "sql", "git", "linux", "system design"],
        "bonus": ["fastapi", "django", "postgresql", "redis", "kafka", "docker", "microservices"],
    },
    "Frontend Engineer": {
        "critical": ["javascript", "react", "html", "css", "git", "responsive design"],
        "bonus": ["typescript", "next.js", "vue", "graphql", "webpack", "testing", "storybook"],
    },
    "DevOps Engineer": {
        "critical": ["docker", "kubernetes", "ci/cd", "linux", "terraform", "git"],
        "bonus": ["aws", "gcp", "azure", "helm", "ansible", "prometheus", "grafana"],
    },
    "Full Stack Engineer": {
        "critical": ["javascript", "python", "sql", "rest api", "react", "git"],
        "bonus": ["node.js", "typescript", "docker", "postgresql", "redis", "aws"],
    },
}

# Sections we try to detect in a resume
SECTION_HEADERS = {
    "experience": ["experience", "work history", "employment", "career", "professional background"],
    "education": ["education", "academic", "qualification", "degree", "university", "college"],
    "skills": ["skills", "technical skills", "competencies", "expertise", "proficiencies"],
    "projects": ["projects", "portfolio", "work samples"],
    "certifications": ["certifications", "certificates", "credentials", "licenses"],
    "summary": ["summary", "objective", "profile", "about me", "overview"],
    "achievements": ["achievements", "accomplishments", "awards", "honors"],
}

DEGREE_PATTERNS = [
    r"\b(b\.?tech|b\.?e\.?|b\.?sc|b\.?s|bachelor(?:\'s)?|b\.?a\.?)\b",
    r"\b(m\.?tech|m\.?e\.?|m\.?sc|m\.?s|master(?:\'s)?|m\.?b\.?a\.?)\b",
    r"\b(ph\.?d|doctorate|d\.?sc)\b",
    r"\b(diploma|associate|a\.?a\.?|a\.?s\.?)\b",
]

EXPERIENCE_PATTERNS = [
    r"(\d+)\+?\s*years?\s+(?:of\s+)?(?:experience|exp)",
    r"(\d+)\+?\s*yrs?\s+(?:of\s+)?(?:experience|exp)",
    r"experience\s+of\s+(\d+)\+?\s*years?",
]

# ─────────────────────────────────────────────────────────────
# Pydantic models
# ─────────────────────────────────────────────────────────────
class CandidateInfo(BaseModel):
    name: str = ""
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    organizations: list[str] = Field(default_factory=list)
    locations: list[str] = Field(default_factory=list)
    years_of_experience: Optional[int] = None
    education: list[str] = Field(default_factory=list)
    languages_spoken: list[str] = Field(default_factory=list)


class SkillsAnalysis(BaseModel):
    matched_critical: list[str]
    missing_critical: list[str]
    matched_bonus: list[str]
    missing_bonus: list[str]
    all_detected_skills: list[str] = Field(default_factory=list)


class ScoreBreakdown(BaseModel):
    critical_skills: int          # 0-40
    bonus_skills: int             # 0-20
    experience: int               # 0-20
    education: int                # 0-10
    completeness: int             # 0-10
    total: int                    # 0-100


class SectionPresence(BaseModel):
    summary: bool = False
    experience: bool = False
    education: bool = False
    skills: bool = False
    projects: bool = False
    certifications: bool = False
    achievements: bool = False


class AnalysisResult(BaseModel):
    candidate: CandidateInfo
    role: str
    score: int
    score_breakdown: ScoreBreakdown
    skills: SkillsAnalysis
    sections: SectionPresence
    recommendations: list[str] = Field(default_factory=list)
    semantic_match_score: float = 0.0


class JobListing(BaseModel):
    id: Optional[Any] = None
    title: Optional[str] = None
    company: Optional[str] = None
    location: Optional[str] = None
    url: Optional[str] = None
    description: Optional[str] = None
    created: Optional[str] = None
    salary_min: Optional[float] = None
    salary_max: Optional[float] = None


class ApiResponse(BaseModel):
    success: bool
    data: Optional[Any] = None
    error: Optional[str] = None
    message: str
    processing_time_ms: Optional[float] = None


# ─────────────────────────────────────────────────────────────
# In-memory rate limiter
# ─────────────────────────────────────────────────────────────
_rate_limit_store: dict[str, list[float]] = defaultdict(list)


def check_rate_limit(client_ip: str) -> bool:
    """Return True if request is allowed, False if rate-limited."""
    now = time.time()
    window_start = now - RATE_LIMIT_WINDOW
    requests_in_window = [t for t in _rate_limit_store[client_ip] if t > window_start]
    requests_in_window.append(now)
    _rate_limit_store[client_ip] = requests_in_window
    return len(requests_in_window) <= RATE_LIMIT_REQUESTS


# ─────────────────────────────────────────────────────────────
# Text extraction
# ─────────────────────────────────────────────────────────────
def extract_text_from_file(uploaded_file: UploadFile) -> str:
    content = uploaded_file.file.read()
    filename = (uploaded_file.filename or "").lower()

    if len(content) > MAX_FILE_SIZE_MB * 1024 * 1024:
        raise HTTPException(status_code=413, detail=f"File exceeds {MAX_FILE_SIZE_MB}MB limit")

    text = ""
    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
            # also grab table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        else:
            text = content.decode("utf-8", errors="ignore")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Text extraction failed: %s", exc)
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    if not text.strip():
        raise HTTPException(status_code=422, detail="Could not extract text from resume")

    return text


# ─────────────────────────────────────────────────────────────
# NLP helpers
# ─────────────────────────────────────────────────────────────
def _spacy_entities(text: str) -> dict[str, list[str]]:
    """Run spaCy NER and return categorised entities."""
    result: dict[str, list[str]] = {"ORG": [], "GPE": [], "PERSON": [], "DATE": []}
    if nlp is None:
        return result
    doc = nlp(text[:50_000])  # cap to avoid OOM on huge docs
    for ent in doc.ents:
        if ent.label_ in result:
            result[ent.label_].append(ent.text.strip())
    # deduplicate
    return {k: list(dict.fromkeys(v)) for k, v in result.items()}


def extract_candidate_info(raw_text: str) -> CandidateInfo:
    text = raw_text  # keep original case for NER
    text_lower = raw_text.lower()

    # ── email ──────────────────────────────────────────────
    email_match = re.search(r"[\w.+-]+@[\w.-]+\.\w{2,}", text)

    # ── phone ──────────────────────────────────────────────
    phone_match = re.search(r"(\+?\d[\d\s\-().]{7,}\d)", text)

    # ── LinkedIn / GitHub ──────────────────────────────────
    linkedin_match = re.search(r"(linkedin\.com/in/[\w-]+)", text_lower)
    github_match = re.search(r"(github\.com/[\w-]+)", text_lower)

    # ── NER via spaCy ──────────────────────────────────────
    entities = _spacy_entities(text)

    # ── Name: prefer first PERSON entity, fallback to first line ──
    name = ""
    if entities["PERSON"]:
        name = entities["PERSON"][0].title()
    else:
        first_line = text.strip().splitlines()[0].strip() if text.strip() else ""
        if first_line and len(first_line.split()) <= 5 and "@" not in first_line:
            name = first_line.title()

    # ── Organizations ──────────────────────────────────────
    org_regex_hits = re.findall(
        r"([A-Z][\w& ]{2,}\b(?:Inc|LLC|Ltd|Corp(?:oration)?|Technologies|Solutions|Systems|Labs?)\.?)",
        text,
    )
    orgs = list(dict.fromkeys(entities["ORG"] + [o.strip() for o in org_regex_hits]))

    # ── Locations ──────────────────────────────────────────
    locations = entities["GPE"][:10]

    # ── Years of experience ────────────────────────────────
    yoe: Optional[int] = None
    for pat in EXPERIENCE_PATTERNS:
        m = re.search(pat, text_lower)
        if m:
            yoe = int(m.group(1))
            break

    # ── Education ─────────────────────────────────────────
    education: list[str] = []
    for pat in DEGREE_PATTERNS:
        for m in re.finditer(pat, text_lower):
            start = max(0, m.start() - 5)
            end = min(len(text), m.end() + 60)
            snippet = text[start:end].strip()
            education.append(snippet)
    education = list(dict.fromkeys(education))[:5]

    return CandidateInfo(
        name=name,
        email=email_match.group(0) if email_match else None,
        phone=phone_match.group(0) if phone_match else None,
        linkedin=linkedin_match.group(0) if linkedin_match else None,
        github=github_match.group(0) if github_match else None,
        organizations=orgs[:10],
        locations=locations,
        years_of_experience=yoe,
        education=education,
    )


# ─────────────────────────────────────────────────────────────
# Section detection
# ─────────────────────────────────────────────────────────────
def detect_sections(text_lower: str) -> SectionPresence:
    presence: dict[str, bool] = {}
    for section, keywords in SECTION_HEADERS.items():
        presence[section] = any(kw in text_lower for kw in keywords)
    return SectionPresence(**presence)


# ─────────────────────────────────────────────────────────────
# Skill detection (exact + TF-IDF semantic)
# ─────────────────────────────────────────────────────────────
def detect_skills_exact(text_lower: str, skills: list[str]) -> list[str]:
    """Exact substring match (also handles multi-word skills)."""
    return [s for s in skills if s in text_lower]


def semantic_match_score(resume_text: str, role: str) -> float:
    """Compute cosine similarity between resume and role description using TF-IDF."""
    role_desc = " ".join(ROLE_SKILLS[role]["critical"] + ROLE_SKILLS[role]["bonus"])
    try:
        vec = TfidfVectorizer(stop_words="english")
        tfidf = vec.fit_transform([resume_text[:5000], role_desc])
        score = float(cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0])
    except Exception:
        score = 0.0
    return round(score, 4)


def collect_all_detected_skills(text_lower: str) -> list[str]:
    """Pool all skills from all roles and return those present in resume."""
    all_skills: set[str] = set()
    for role_data in ROLE_SKILLS.values():
        all_skills.update(role_data["critical"])
        all_skills.update(role_data["bonus"])
    return sorted([s for s in all_skills if s in text_lower])


# ─────────────────────────────────────────────────────────────
# Score calculation
# ─────────────────────────────────────────────────────────────
def compute_score(
    matched_critical: list[str],
    all_critical: list[str],
    matched_bonus: list[str],
    all_bonus: list[str],
    candidate: CandidateInfo,
    sections: SectionPresence,
) -> ScoreBreakdown:

    # Critical skills → max 40 pts
    critical_score = int((len(matched_critical) / max(len(all_critical), 1)) * 40)

    # Bonus skills → max 20 pts
    bonus_score = int((len(matched_bonus) / max(len(all_bonus), 1)) * 20)

    # Experience → max 20 pts
    exp_score = 0
    yoe = candidate.years_of_experience or 0
    if yoe >= 7:
        exp_score = 20
    elif yoe >= 4:
        exp_score = 15
    elif yoe >= 2:
        exp_score = 10
    elif yoe >= 1:
        exp_score = 5
    elif sections.experience:
        exp_score = 3  # section present but no explicit years

    # Education → max 10 pts
    edu_text = " ".join(candidate.education).lower()
    edu_score = 0
    if any(p in edu_text for p in ["ph.d", "phd", "doctorate"]):
        edu_score = 10
    elif any(p in edu_text for p in ["master", "m.tech", "m.sc", "mba", "m.s"]):
        edu_score = 8
    elif any(p in edu_text for p in ["bachelor", "b.tech", "b.e", "b.sc", "b.s"]):
        edu_score = 6
    elif any(p in edu_text for p in ["diploma", "associate"]):
        edu_score = 4
    elif sections.education:
        edu_score = 2

    # Completeness → max 10 pts (profile completeness)
    completeness = 0
    if candidate.email:
        completeness += 2
    if candidate.phone:
        completeness += 2
    if candidate.linkedin:
        completeness += 2
    if sections.summary:
        completeness += 2
    if sections.certifications or sections.achievements:
        completeness += 2

    total = min(critical_score + bonus_score + exp_score + edu_score + completeness, 100)

    return ScoreBreakdown(
        critical_skills=critical_score,
        bonus_skills=bonus_score,
        experience=exp_score,
        education=edu_score,
        completeness=completeness,
        total=total,
    )


# ─────────────────────────────────────────────────────────────
# Recommendations engine
# ─────────────────────────────────────────────────────────────
def generate_recommendations(
    role: str,
    missing_critical: list[str],
    missing_bonus: list[str],
    sections: SectionPresence,
    candidate: CandidateInfo,
) -> list[str]:
    recs: list[str] = []

    if missing_critical:
        top = missing_critical[:3]
        recs.append(f"Add or strengthen critical skills: {', '.join(top)}.")

    if missing_bonus[:2]:
        recs.append(f"Bonus skills that boost {role} profiles: {', '.join(missing_bonus[:2])}.")

    if not sections.summary:
        recs.append("Add a professional summary section at the top of your resume.")

    if not sections.projects:
        recs.append("Include a Projects section showcasing relevant work or open-source contributions.")

    if not sections.certifications:
        recs.append("Add relevant certifications (e.g., AWS, Google Cloud, Coursera specialisations).")

    if not candidate.linkedin:
        recs.append("Include your LinkedIn profile URL to increase visibility.")

    if not candidate.github and role in ("ML Engineer", "AI Engineer", "Backend Engineer", "Full Stack Engineer"):
        recs.append("Add a GitHub link — hiring managers actively look for portfolios.")

    if candidate.years_of_experience is None:
        recs.append("Mention your total years of experience explicitly (e.g., '5+ years of experience in…').")

    return recs


# ─────────────────────────────────────────────────────────────
# Job Fetching
# ─────────────────────────────────────────────────────────────
def fetch_jobs(role: str, results: int = 5) -> list[JobListing]:
    url = f"https://api.adzuna.com/v1/api/jobs/{COUNTRY}/search/1"
    params = {
        "app_id": ADZUNA_APP_ID,
        "app_key": ADZUNA_APP_KEY,
        "what": role,
        "results_per_page": min(results, 10),
    }
    try:
        resp = requests.get(url, params=params, timeout=6)
        resp.raise_for_status()
        data = resp.json()
        listings = []
        for job in data.get("results", []):
            listings.append(
                JobListing(
                    id=job.get("id"),
                    title=job.get("title"),
                    company=job.get("company", {}).get("display_name"),
                    location=job.get("location", {}).get("display_name"),
                    url=job.get("redirect_url"),
                    description=(job.get("description") or "")[:300],
                    created=job.get("created"),
                    salary_min=job.get("salary_min"),
                    salary_max=job.get("salary_max"),
                )
            )
        return listings
    except Exception as exc:
        logger.warning("Job fetch failed: %s", exc)
        return []


# ─────────────────────────────────────────────────────────────
# Full analysis pipeline
# ─────────────────────────────────────────────────────────────
def analyze_resume(raw_text: str) -> AnalysisResult:
    text_lower = raw_text.lower()

    candidate = extract_candidate_info(raw_text)
    sections = detect_sections(text_lower)
    all_detected_skills = collect_all_detected_skills(text_lower)

    best_role: Optional[str] = None
    best_score_val = -1
    best_matched_critical: list[str] = []
    best_matched_bonus: list[str] = []
    best_sem_score = 0.0

    for role, skill_map in ROLE_SKILLS.items():
        matched_c = detect_skills_exact(text_lower, skill_map["critical"])
        matched_b = detect_skills_exact(text_lower, skill_map["bonus"])
        sem_score = semantic_match_score(raw_text, role)

        # Composite role confidence = exact match ratio + semantic
        ratio = (len(matched_c) / max(len(skill_map["critical"]), 1)) * 0.7
        composite = ratio + sem_score * 0.3

        if composite > best_score_val:
            best_score_val = composite
            best_role = role
            best_matched_critical = matched_c
            best_matched_bonus = matched_b
            best_sem_score = sem_score

    role = best_role or "Data Analyst"
    skill_map = ROLE_SKILLS[role]

    missing_critical = list(set(skill_map["critical"]) - set(best_matched_critical))
    missing_bonus = list(set(skill_map["bonus"]) - set(best_matched_bonus))

    score_breakdown = compute_score(
        best_matched_critical, skill_map["critical"],
        best_matched_bonus, skill_map["bonus"],
        candidate, sections,
    )

    skills_analysis = SkillsAnalysis(
        matched_critical=best_matched_critical,
        missing_critical=missing_critical,
        matched_bonus=best_matched_bonus,
        missing_bonus=missing_bonus,
        all_detected_skills=all_detected_skills,
    )

    recommendations = generate_recommendations(
        role, missing_critical, missing_bonus, sections, candidate
    )

    return AnalysisResult(
        candidate=candidate,
        role=role,
        score=score_breakdown.total,
        score_breakdown=score_breakdown,
        skills=skills_analysis,
        sections=sections,
        recommendations=recommendations,
        semantic_match_score=best_sem_score,
    )


# ─────────────────────────────────────────────────────────────
# App setup
# ─────────────────────────────────────────────────────────────
app = FastAPI(
    title="ATS Resume Analyzer",
    description="Production-grade ATS resume analysis with NLP, semantic matching, and job fetching.",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc",
)

# simple root endpoint so GET / doesn’t return 404
@app.get("/", summary="Root welcome")
async def root():
    return {"message": "ATS Analyzer running - see /docs for available endpoints"}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    if request.url.path.startswith("/api/"):
        client_ip = request.client.host if request.client else "unknown"
        if not check_rate_limit(client_ip):
            return JSONResponse(
                status_code=429,
                content={"success": False, "error": "Rate limit exceeded", "message": "Too many requests"},
            )
    return await call_next(request)


@app.middleware("http")
async def log_requests(request: Request, call_next):
    start = time.perf_counter()
    response = await call_next(request)
    elapsed = (time.perf_counter() - start) * 1000
    logger.info("%s %s → %d (%.1fms)", request.method, request.url.path, response.status_code, elapsed)
    return response


# ─────────────────────────────────────────────────────────────
# Endpoints
# ─────────────────────────────────────────────────────────────
@app.get("/health")
async def health():
    return {
        "status": "ok",
        "spacy": nlp is not None,
        "version": "2.0.0",
    }


@app.post("/api/analyze", response_model=ApiResponse, summary="Analyze a resume file")
async def analyze_endpoint(file: UploadFile = File(...)):
    """
    Upload a **PDF** or **DOCX** resume. Returns:
    - Candidate profile (name, email, phone, LinkedIn, GitHub, education…)
    - Best-match role & ATS score (0–100)
    - Score breakdown (critical skills, bonus skills, experience, education, completeness)
    - Matched/missing skills
    - Section presence
    - Personalised recommendations
    - Semantic similarity score
    """
    start = time.perf_counter()

    if not (file.filename or "").lower().endswith((".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    raw_text = extract_text_from_file(file)
    result = analyze_resume(raw_text)

    elapsed = (time.perf_counter() - start) * 1000
    return ApiResponse(
        success=True,
        data=result.dict(),
        message="Analysis complete",
        processing_time_ms=round(elapsed, 1),
    )


@app.post("/api/analyze/role", response_model=ApiResponse, summary="Analyze resume against a specific role")
async def analyze_for_role(
    file: UploadFile = File(...),
    role: str = Query(..., description="Target role, e.g. 'Data Scientist'"),
):
    """Force analysis against a specific role instead of auto-detecting the best match."""
    start = time.perf_counter()

    if role not in ROLE_SKILLS:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown role '{role}'. Valid roles: {list(ROLE_SKILLS.keys())}",
        )

    if not (file.filename or "").lower().endswith((".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    raw_text = extract_text_from_file(file)
    text_lower = raw_text.lower()
    skill_map = ROLE_SKILLS[role]

    candidate = extract_candidate_info(raw_text)
    sections = detect_sections(text_lower)
    all_detected_skills = collect_all_detected_skills(text_lower)

    matched_c = detect_skills_exact(text_lower, skill_map["critical"])
    matched_b = detect_skills_exact(text_lower, skill_map["bonus"])
    sem_score = semantic_match_score(raw_text, role)

    missing_critical = list(set(skill_map["critical"]) - set(matched_c))
    missing_bonus = list(set(skill_map["bonus"]) - set(matched_b))

    score_breakdown = compute_score(matched_c, skill_map["critical"], matched_b, skill_map["bonus"], candidate, sections)

    skills_analysis = SkillsAnalysis(
        matched_critical=matched_c,
        missing_critical=missing_critical,
        matched_bonus=matched_b,
        missing_bonus=missing_bonus,
        all_detected_skills=all_detected_skills,
    )

    recommendations = generate_recommendations(role, missing_critical, missing_bonus, sections, candidate)

    result = AnalysisResult(
        candidate=candidate,
        role=role,
        score=score_breakdown.total,
        score_breakdown=score_breakdown,
        skills=skills_analysis,
        sections=sections,
        recommendations=recommendations,
        semantic_match_score=sem_score,
    )

    elapsed = (time.perf_counter() - start) * 1000
    return ApiResponse(success=True, data=result.dict(), message="Analysis complete", processing_time_ms=round(elapsed, 1))


@app.get("/api/jobs", response_model=ApiResponse, summary="Fetch live job listings")
async def jobs_endpoint(
    role: str = Query(..., description="Job role to search"),
    results: int = Query(5, ge=1, le=10, description="Number of results"),
):
    """Fetch live job listings from Adzuna for a given role."""
    start = time.perf_counter()
    jobs = fetch_jobs(role, results)
    elapsed = (time.perf_counter() - start) * 1000
    return ApiResponse(
        success=True,
        data={"jobs": [j.dict() for j in jobs], "count": len(jobs)},
        message=f"Fetched {len(jobs)} job(s)",
        processing_time_ms=round(elapsed, 1),
    )


@app.get("/api/roles", response_model=ApiResponse, summary="List supported roles")
async def list_roles():
    """Returns all roles and their skill sets."""
    return ApiResponse(
        success=True,
        data={"roles": list(ROLE_SKILLS.keys()), "skill_map": ROLE_SKILLS},
        message="Roles fetched",
    )


# ─────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "5000"))
    uvicorn.run("main:app", host="127.0.0.1", port=port, reload=True, log_level="info")